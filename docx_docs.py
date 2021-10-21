import os
import docx
import tf_fio
import xgboost as xgb



text_to_replace = 'Noname'
black_list_sym = '1234567890!№"#$%&()*+,-/:;<=>?@[\\]^_`{|}~\t\n'

class Docs():
    def __init__(self, filename):
        self.filename = filename
        self.pagelist = []
        self.directory = '/files/'
        self.output_fname = ''
        self.doc = docx.Document(self.filename)
        self.model = tf_fio.Predictor()


    def table_process(self):

        print('Обрабатываем таблицы...')
        for p in self.doc.tables:
            for t in p.table.rows:
                for c in t.cells:
                    splits = c.text.split(' ')
                    predict = self.model.bst_predict(splits)
                    print(splits)
                    for i, p in enumerate(predict):

                        if (p == 1) and (splits[i] != '') and (splits[i] not in black_list_sym):
                            splits[i] = text_to_replace
                    str = ' '.join(splits)

                    c.text = str



    def paragraph_process(self):

        print('Обрабатываем абзацы...')
        for p in self.doc.paragraphs:
            for run in p.runs:
                splits = run.text.split(' ')
                predict = self.model.bst_predict(splits)
                for i, p in enumerate(predict):
                    if (p == 1) and (splits[i]!='') and (splits[i] not in black_list_sym):
                        splits[i] =text_to_replace
                str = ' '.join(splits)
                run.text = str



    def recogn_doc(self):
        self.paragraph_process()
        self.table_process()
        self.output_fname = os.curdir + self.directory+ 'out_' + self.filename.split("/")[-1]
        self.doc.save(self.output_fname)
        return self.output_fname



